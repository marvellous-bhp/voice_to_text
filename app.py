from flask import Flask, request, jsonify
from llm.llm import *
from voice.Wav2Vec import *
from langchain.schema import Document
from docx import Document as DocxDocument
import os
app = Flask(__name__)
doc_path = "activities house.docx"
docx_document = DocxDocument(doc_path)
docs = [Document(page_content=paragraph.text) for paragraph in docx_document.paragraphs if paragraph.text.strip()]
embeddings=get_embedding_Transformer(model_name_Transformer_384)
retriever=vector_data_chroma(docs,embeddings)
compression_retriever=create_compression_retriever_CohereRerank(retriever)
qa=create_AI_agent(llm_groq,compression_retriever)

UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Function to process voice to text
def process_voice_file(file_path):
    # Sử dụng mô hình xử lý giọng nói của bạn để chuyển đổi từ giọng nói sang văn bản
    text = get_result(file_path, qa)
    return text

# Flask routes
@app.route("/") 
def read_root():
    return {"message": "Hello from Flask"}


@app.route("/status")
def status():
    return {"status": "Service is running!"}

link_voice_test="data/tat_quat (1).wav"
@app.route("/voice2text/<link_voice>")
def voice2text(link_voice):
    result = get_result(link_voice,qa)
    return result

@app.route("/upload", methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    
    if file:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)
        result = process_voice_file(file_path)
        return jsonify({"transcript": result})

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))

    # Run the Flask app
    app.run(host='0.0.0.0', port=port)